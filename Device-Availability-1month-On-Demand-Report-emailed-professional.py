import requests
import os
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
import time
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email.mime.text import MIMEText
from email.utils import COMMASPACE
from email import encoders
from datetime import datetime, timedelta
import calendar
from dotenv import load_dotenv

# Put the recipient_email to receive the report. 
recipient_email = "alexiscool@gmail.com"

# Get the current date
now = datetime.now()

# Calculate the first and last day of the last month
if now.month == 1:
    last_month = 12
    year = now.year - 1
else:
    last_month = now.month - 1
    year = now.year

_, last_day_of_last_month = calendar.monthrange(year, last_month)

first_day_of_last_month = datetime(year, last_month, 1)
last_day_of_last_month = datetime(year, last_month, last_day_of_last_month)

# Convert datetime to UNIX timestamp (seconds since the epoch)
start_timestamp_last_month = int(first_day_of_last_month.timestamp())
end_timestamp_last_month = int(last_day_of_last_month.timestamp())

# Calculate the duration
duration_last_month = end_timestamp_last_month - start_timestamp_last_month

# Now use start_timestamp_last_month, end_timestamp_last_month, and duration_last_month in your code


# Load the environment variables from the .env file
load_dotenv()

# Load sensitive information from environment variables
email_sender = os.environ.get("EMAIL_SENDER")
email_password = os.environ.get("EMAIL_PASSWORD")
api_token = os.environ.get("API_TOKEN")

timestamp = datetime.now().strftime('%Y-%m-%d')

def send_email(subject, body, to, file_path):
    sender = email_sender
    password = email_password

    msg = MIMEMultipart()
    msg['From'] = sender
    msg['To'] = COMMASPACE.join(to)
    msg['Subject'] = subject

    msg.attach(MIMEText(body))

    with open(file_path, "rb") as f:
        part = MIMEBase('application', 'octet-stream')
        part.set_payload(f.read())
        encoders.encode_base64(part)
        part.add_header('Content-Disposition', f'attachment; filename="{os.path.basename(file_path)}"')
        msg.attach(part)

    server = smtplib.SMTP_SSL('smtp.gmail.com', 465)
    server.ehlo()
    server.login(sender, password)
    server.sendmail(sender, to, msg.as_string())
    server.close()

device_groups = ["C2 VAL HE03-SG03101 L-CHTR QA02 Canary Monitoring", "C2 VAL HE11-SG11201 L-CHTR QA02 Canary Monitoring", "C2 VAL HE20-SG20101 L-CHTR UAT Canary Monitoring", "C2 VAL HE23-H1-SG1 L-CHTR UAT Canary Monitoring", "C2 VAL HE24-SG24101 L-CHTR UAT Canary Monitoring", "SP HE03-SG03101 L-CHTR QA02 Canary Monitoring", "SP HE20-SG20102 L-CHTR UAT Canary Monitoring", "C2 VAL HE02-SG02118 L-TWC UAT Canary Monitoring", "C2 VAL HE05-SG05101 L-TWC QA02 Canary Monitoring"]


base_url = "http://172.30.121.19/api/v0"
headers = {'X-Auth-Token': api_token}

def get_device_availability(device_id):
    url = f"{base_url}/devices/{device_id}/availability"
    response = requests.get(url, headers=headers)
    return response.json()

def download_availability_graph(device_id, duration, start_timestamp, end_timestamp):
    graph_url = f"http://172.30.121.19/graph.php?device={device_id}&type=device_availability&duration={duration}&from={start_timestamp}&to={end_timestamp}&height=309&width=1134.9"
    response = requests.get(graph_url, headers=headers, stream=True)

    # Create the "graph" folder if it doesn't exist
    if not os.path.exists("graph"):
        os.makedirs("graph")

    file_path = f"graph/device_{device_id}_availability_{duration}.png"

    if response.status_code == 200:
        with open(file_path, "wb") as img_file:
            for chunk in response.iter_content(1024):
                img_file.write(chunk)
        return file_path
    else:
        print(f"Error downloading graph for device {device_id} with duration {duration}: {response.status_code}")
        return None

def format_availability(availability):
    formatted_availability = []
    for item in availability:
        duration = item['duration']
        availability_perc = item['availability_perc']
        if duration == 86400:
            duration_text = '1 day'
        elif duration == 604800:
            duration_text = '1 week'
        elif duration == 2592000:
            duration_text = '1 month'
        elif duration == 31536000:
            duration_text = '1 year'
        else:
            duration_text = f'{duration} seconds'
        formatted_availability.append(f'{duration_text}: {availability_perc}%')
    return ', '.join(formatted_availability)

document = Document()

# Define styles for headings and paragraphs
styles = document.styles

heading1_style = styles.add_style('CustomHeading1', WD_STYLE_TYPE.PARAGRAPH)
heading1_style.base_style = styles['Heading 1']
heading1_style.font.name = 'Arial'
heading1_style.font.size = Pt(16)
heading1_style.font.bold = True

heading2_style = styles.add_style('CustomHeading2', WD_STYLE_TYPE.PARAGRAPH)
heading2_style.base_style = styles['Heading 2']
heading2_style.font.name = 'Arial'
heading2_style.font.size = Pt(14)
heading2_style.font.bold = True

normal_style = styles.add_style('CustomNormal', WD_STYLE_TYPE.PARAGRAPH)
normal_style.base_style = styles['Normal']
normal_style.font.name = 'Arial'
normal_style.font.size = Pt(12)

document.add_heading(f'Monthly Modem Availability Report {timestamp}', 0).style = heading1_style

for device_group in device_groups:
    document.add_heading(device_group, level=1).style = heading2_style  # Add the device group label to the report

    librenms_url = f"{base_url}/devicegroups/{device_group}"
    librenms_devices = requests.get(librenms_url, headers=headers).json()

    # Calculate the average availability for the device group
    total_availability = 0
    num_devices = 0
    for device in librenms_devices["devices"]:
        device_id = device["device_id"]

        # Get the availability data and format it
        availability_data = get_device_availability(device_id)

        # Filter the availability data to only include data for the last week
        availability_data_last_week = [item for item in availability_data['availability'] if item['duration'] == 2592000]

        if availability_data_last_week:
            total_availability += float(availability_data_last_week[0]['availability_perc'])
            num_devices += 1

    average_availability = total_availability / num_devices if num_devices > 0 else 0
    document.add_paragraph(f"Average availability for {device_group}: {average_availability:.2f}%").style = normal_style

    # Add a table to the document
    table = document.add_table(rows=1, cols=3)
    table.style = 'Table Grid'

    # Add header row
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Device Name'
    header_cells[1].text = 'Availability'
    header_cells[2].text = 'Graph'

    # Add data rows
    for device in librenms_devices["devices"]:
        device_id = device["device_id"]
        
        # Retrieve the device information from the API
        device_url = f"{base_url}/devices/{device_id}"
        device_info = requests.get(device_url, headers=headers).json()
        
        # Retrieve the device display name from the API response
        # Use a default value (the device ID) if the "display" field is not present
        device_name = device_info["devices"][0].get("display", str(device_id))
        
        availability_data = get_device_availability(device_id)
        availability_data_last_week = [item for item in availability_data['availability'] if item['duration'] == 2592000]
        formatted_availability = format_availability(availability_data_last_week)

        row_cells = table.add_row().cells
        row_cells[0].text = device_name # Use the device display name instead of the device ID
        row_cells[1].text = formatted_availability

        end_timestamp = int(time.time())
        start_timestamp = end_timestamp - 2592000 
        duration = 2592000
        graph_path = download_availability_graph(device_id, duration, start_timestamp, end_timestamp)
        if graph_path:
            row_cells[2].paragraphs[0].add_run().add_picture(graph_path, width=Inches(2))

report_filename = f"Monthly_Modem_Availability_Report_{timestamp}.docx"
document.save(report_filename)

subject = f"Monthly Modem Availability Report {timestamp}"
body = f"Please find attached the Monthly Modem Availability Report for {timestamp}."
send_email(subject, body, recipient_email, report_filename)

print("Report sent.")
